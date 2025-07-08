import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook


def ensure_data_dir_exists():
    """Создает папку 'data', если её нет."""
    os.makedirs("data", exist_ok=True)


def save_to_txt(data_dict, filename):
    """Сохраняет данные в TXT-файл."""
    filepath = os.path.join("data", f"{filename}.txt")
    with open(filepath, "w", encoding="utf-8") as f:
        for key, value in data_dict.items():
            if value:
                f.write(f'{key}: {value}\n')
    return filepath


def save_to_docx(data_dict, filename):
    """Сохраняет данные в DOCX (Word)."""
    filepath = os.path.join("data", f"{filename}.docx")
    doc = Document()
    doc.add_heading("Данные системы", level=1)
    for key, value in data_dict.items():
        if value:
            doc.add_paragraph(f"{key}: {value}")
    doc.save(filepath)
    return filepath


def save_to_xlsx(data_dict, filename):
    """Сохраняет данные в XLSX (Excel)."""
    filepath = os.path.join("data", f"{filename}.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "System Data"
    ws.append(["Характеристики", "Значения"])
    for key, value in data_dict.items():
        if value:
            ws.append([key, str(value)])
    wb.save(filepath)
    return filepath


def save_to_html(data_dict, filename):
    """Сохраняет данные в HTML."""
    filepath = os.path.join("data", f"{filename}.html")
    html_content = """
    <html>
        <head><meta charset="UTF-8"></head>
        <body>
            <h1>Данные системы</h1>
            <table border="1">
                <tr><th>Характеристики</th><th>Значения</th></tr>
    """
    for key, value in data_dict.items():
        if value:
            html_content += f"<tr><td><b>{key}</b></td><td>{value}</td></tr>"
    html_content += "</table></body></html>"

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)
    return filepath


def save_data_to_all_formats(data_dict, filename="output"):
    """Сохраняет данные во всех доступных форматах."""
    ensure_data_dir_exists()

    saved_files = {
        "TXT": save_to_txt(data_dict, filename),
        "DOCX": save_to_docx(data_dict, filename),
        "XLSX": save_to_xlsx(data_dict, filename),
        "HTML": save_to_html(data_dict, filename),
    }

    print("✅ Файлы сохранены в папку 'data':")
    for format_name, path in saved_files.items():
        if path:  # Если файл успешно создан
            print(f"  - {format_name}: {os.path.basename(path)}")

    return saved_files


if __name__ == "__main__":
    system_data = {
        "OS": "Windows 11",
        "CPU": "AMD Ryzen 7",
        "RAM": "32 GB",
        "Диск": "1 TB NVMe SSD",
        "123": None
    }
    current_datetime = datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
    save_data_to_all_formats(system_data, f"system_report_{current_datetime}")